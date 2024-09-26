import logging

from docx import Document
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

from src.module2 import generate_script

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_document(model, user_input_prompt, content):
    prompt = f"""I will give you an input. If it's related to the business domain, make an outline with paragraphs to describe this product. Each paragraph must contain at least 5 sentences! If it's related to the education domain, make an outline with paragraphs to create a comprehensive lesson plan for teaching the given topic. Each paragraph must contain at least 5 sentences! Ensure each outline is clear, structured, and focuses on key elements of the domain (product features for business, learning objectives for education), without including additional commentary. Each paragraph must contain at least 5 sentences! The output must strictly be a JSON object, formatted as follows:
    \"\"\"
    {{
    "outline": [
    {{"heading for idea 1": ["paragraph 1 of idea 1", "paragraph 2 of idea 1",...]}},
    {{"heading for idea 2": ["paragraph 1 of idea 2", "paragraph 2 of idea 2",...]}}
    ...
    ]
    }}
    \"\"\"

    And the input is:
    \"\"\"
    {user_input_prompt}
    \"\"\"
    The output must match this structure exactly. MAKE SURE THAT LANGUAGE OF USER INPUT IS WHAT LANGUAGE YOU USE IN THE RESPONSE. Each paragraph must contain at least 5 sentences! Each paragraph must contain at least 5 sentences! Each paragraph must contain at least 5 sentences! Each paragraph must contain at least 5 sentences! Each paragraph must contain at least 5 sentences! Each paragraph must contain at least 5 sentences! Each paragraph must contain at least 5 sentences!"""
    outline_object = generate_script(model, prompt)
    print(outline_object)
    keys = []
    result_json = {}
    doc = Document()
    for element in outline_object['outline']:
        key = list(element.keys())[0]
        keys.append(key)
        doc.add_heading(key, level=1)  # Add Heading
        for idea in element[key]:
            doc.add_paragraph(idea)
            
    # Save the document to a file
    doc.save("./temp/output.docx")

    logger.info("Document created successfully!")
    return "./temp/output.docx"


if __name__ == "__main__":
    genai.configure(api_key="")
    model = genai.GenerativeModel(model_name='gemini-1.5-flash')
    create_document(model, "Cho tôi một video để nói về sản phẩm Bia Hơi Hà Nội - Thùng 24 lon 500ml - Phiên bản Tết", "non")
    